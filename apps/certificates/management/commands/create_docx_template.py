from django.core.management.base import BaseCommand
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from django.conf import settings


class Command(BaseCommand):
    help = "Generates a starter DOCX template for certificates"

    def handle(self, *args, **options):
        # Ensure media directory exists
        template_dir = os.path.join(settings.MEDIA_ROOT, "templates")
        os.makedirs(template_dir, exist_ok=True)

        doc = Document()

        # Set margins suitable for A4
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("Republic of the Philippines\n")
        run.font.name = "Arial"
        run.font.size = Pt(12)

        run = header.add_run("Province of {{ province }}\n")
        run.font.name = "Arial"
        run.font.size = Pt(12)

        run = header.add_run("City/Municipality of {{ city }}\n")
        run.font.name = "Arial"
        run.font.size = Pt(12)

        run = header.add_run("BARANGAY {{ barangay_name }}\n")
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.bold = True

        run = header.add_run("OFFICE OF THE PUNONG BARANGAY\n")
        run.font.name = "Arial"
        run.font.size = Pt(16)
        run.font.bold = True

        doc.add_paragraph()  # Spacer

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("{{ certificate_title }}")
        run.font.name = "Arial"
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.underline = True

        doc.add_paragraph()  # Spacer

        # Salutation
        p = doc.add_paragraph()
        run = p.add_run("TO WHOM IT MAY CONCERN:")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.bold = True

        doc.add_paragraph()  # Spacer

        # Body
        body = doc.add_paragraph()
        body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = body.add_run("    This is to certify that ")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        run = body.add_run("{{ full_name }}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.bold = True

        run = body.add_run(", of legal age, {{ citizenship }}, and a resident of ")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        run = body.add_run("{{ address }}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.italic = True

        run = body.add_run(", has ")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        run = body.add_run("NO DEROGATORY RECORD")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.bold = True

        run = body.add_run(" on file with this office as of this date.")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        doc.add_paragraph()  # Spacer

        # Purpose
        purpose = doc.add_paragraph()
        purpose.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = purpose.add_run(
            "    This certification is being issued upon the request of the above-mentioned person for "
        )
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        run = purpose.add_run("{{ purpose }}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.bold = True

        run = purpose.add_run(" and for whatever legal purpose it may serve.")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        doc.add_paragraph()  # Spacer
        doc.add_paragraph()  # Spacer

        # Date
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(
            "Issued this {{ day }} day of {{ month_year }} at the Office of the Punong Barangay."
        )
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.italic = True

        doc.add_paragraph()  # Spacer
        doc.add_paragraph()  # Spacer
        doc.add_paragraph()  # Spacer

        # Signature
        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = sig.add_run("HON. RICARDO VELASCO\n")
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.bold = True

        run = sig.add_run("Punong Barangay")
        run.font.name = "Arial"
        run.font.size = Pt(11)

        # Save
        filename = "certificate_template.docx"
        filepath = os.path.join(template_dir, filename)
        doc.save(filepath)

        self.stdout.write(
            self.style.SUCCESS(f"Successfully created starter template: {filepath}")
        )
